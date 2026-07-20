from __future__ import annotations

import os
import shutil
import tempfile
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from .ats_cv import (
    ATS_TEMPLATE_PATH,
    ATS_VARIANTS,
    build_all_ats_cvs,
)
from .core import load_json, sha256_file
from .cv_composer import compose_standard_cv
from .cv_model import CVViewModel
from .pdf_renderer import PAGE_HEIGHT, PAGE_WIDTH, render_standard_cv

STANDARD_CV_FILENAME = "cv-delia-rossignol-signature.pdf"
TEMPLATE_PATH = Path("templates/cv/signature-editorial/template.json")
STRATEGY_PATH = Path("data/style/cv-standard.json")
REQUIRED_CV_TEXT = ("construire ensemble une solution concrète",)
FORBIDDEN_CV_TEXT = ("41 ans", "2 enfants", "30 octobre 1978", "5, rue Jacques Offenbach")


@dataclass
class DocumentCheckState:
    root: Path
    published: Path
    template: dict[str, Any]
    view: CVViewModel
    errors: list[str] = field(default_factory=list)
    first_result: dict[str, Any] = field(default_factory=dict)

    def report(self) -> dict[str, Any]:
        return {
            "documents": [self.first_result],
            "errors": self.errors,
            "ok": not self.errors,
            "published_output": str(self.published.resolve()),
        }


def _inputs(root: Path) -> tuple[dict[str, Any], dict[str, Any]]:
    return load_json(root / TEMPLATE_PATH), load_json(root / STRATEGY_PATH)


def _build_standard_cv_from_view(template: dict[str, Any], view: CVViewModel, output: Path) -> dict[str, Any]:
    layout = render_standard_cv(view, output)
    return {
        "id": "standard-cv-signature-editorial",
        "output": str(output.resolve()),
        "pages": int(template["rendering"]["preferred_pages"]),
        "sha256": sha256_file(output),
        "template": view.template_id,
        "template_version": view.template_version,
        "content_strategy": view.strategy_id,
        "source_ids": list(view.source_ids),
        "layout": layout,
    }


def build_standard_cv(root: Path, output: Path) -> dict[str, Any]:
    root = root.resolve()
    output = output if output.is_absolute() else root / output
    template, strategy = _inputs(root)
    view = compose_standard_cv(root, template, strategy)
    return _build_standard_cv_from_view(template, view, output)


def _atomic_copy(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.{uuid.uuid4().hex}.tmp")
    try:
        shutil.copyfile(source, temporary)
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)


def build_documents(
    root: Path,
    output_dir: Path | None = None,
    public_dir: Path | None = None,
) -> dict[str, Any]:
    root = root.resolve()
    output_dir = output_dir or root / "output" / "pdf"
    public_dir = public_dir or root / "site" / "assets" / "downloads"
    if not output_dir.is_absolute():
        output_dir = root / output_dir
    if not public_dir.is_absolute():
        public_dir = root / public_dir
    output = output_dir / STANDARD_CV_FILENAME
    public_asset = public_dir / STANDARD_CV_FILENAME
    result = build_standard_cv(root, output)
    _atomic_copy(output, public_asset)
    result["public_output"] = str(public_asset.resolve())
    ats_results = build_all_ats_cvs(root, output_dir)
    for ats_result in ats_results:
        source = Path(str(ats_result["output"]))
        published = public_dir / source.name
        _atomic_copy(source, published)
        ats_result["public_output"] = str(published.resolve())
    return {"documents": [result, *ats_results], "ok": True}


def _generate_check_copies(state: DocumentCheckState, temporary: Path) -> Path:
    first = temporary / "first.pdf"
    second = temporary / "second.pdf"
    state.first_result = _build_standard_cv_from_view(state.template, state.view, first)
    _build_standard_cv_from_view(state.template, state.view, second)
    first_bytes = first.read_bytes()
    if first_bytes != second.read_bytes():
        state.errors.append("standard CV generation is not byte-for-byte reproducible")
    if not state.published.is_file():
        state.errors.append(f"published CV is missing: {state.published.relative_to(state.root)}")
    elif first_bytes != state.published.read_bytes():
        state.errors.append(f"published CV is stale: {state.published.relative_to(state.root)}")
    return first


def _check_pdf_pages(state: DocumentCheckState, reader: PdfReader) -> None:
    maximum_pages = int(state.template["rendering"]["maximum_pages"])
    preferred_pages = int(state.template["rendering"]["preferred_pages"])
    if len(reader.pages) != preferred_pages:
        state.errors.append(f"standard CV must contain {preferred_pages} pages, found {len(reader.pages)}")
    if len(reader.pages) > maximum_pages:
        state.errors.append(f"standard CV exceeds maximum page count: {maximum_pages}")
    for index, page in enumerate(reader.pages, start=1):
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        if abs(width - PAGE_WIDTH) > 0.1 or abs(height - PAGE_HEIGHT) > 0.1:
            state.errors.append(f"page {index} is not A4: {width:.2f} x {height:.2f} points")


def _check_pdf_text(state: DocumentCheckState, reader: PdfReader) -> None:
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    for required in (state.view.email, state.view.phone, *REQUIRED_CV_TEXT):
        if required not in text:
            state.errors.append(f"standard CV is missing required text: {required}")
    for forbidden in FORBIDDEN_CV_TEXT:
        if forbidden.casefold() in text.casefold():
            state.errors.append(f"standard CV contains forbidden text: {forbidden}")


def _check_rendered_layout(state: DocumentCheckState) -> None:
    layout = state.first_result.get("layout", {})
    if float(layout.get("minimum_y", 0)) < 40:
        state.errors.append("standard CV content crosses the safe lower page boundary")
    audit = layout.get("audit", {})
    state.errors.extend(f"standard CV layout violation: {violation}" for violation in audit.get("violations", []))
    if int(audit.get("overflow_count", 0)):
        state.errors.append("standard CV layout contains overflowing graphical elements")


def _document_text(document: Any) -> str:
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _check_ats_reproducibility(
    root: Path,
    published_dir: Path,
    temporary: Path,
) -> tuple[list[dict[str, Any]], list[str]]:
    errors: list[str] = []
    first_dir = temporary / "ats-first"
    second_dir = temporary / "ats-second"
    first_results = build_all_ats_cvs(root, first_dir)
    build_all_ats_cvs(root, second_dir)
    for variant in ATS_VARIANTS:
        for filename in (variant.pdf_filename, variant.docx_filename):
            first = first_dir / filename
            second = second_dir / filename
            if first.read_bytes() != second.read_bytes():
                errors.append(f"ATS CV generation is not byte-for-byte reproducible: {filename}")
            published = published_dir / filename
            if not published.is_file():
                errors.append(f"published ATS CV is missing: {published.relative_to(root)}")
            elif first.read_bytes() != published.read_bytes():
                errors.append(f"published ATS CV is stale: {published.relative_to(root)}")
    return first_results, errors


def _check_ats_pdf(root: Path, path: Path, required_keywords: tuple[str, ...]) -> list[str]:
    errors: list[str] = []
    template = load_json(root / ATS_TEMPLATE_PATH)
    reader = PdfReader(str(path))
    maximum_pages = int(template["rendering"]["maximum_pages"])
    if len(reader.pages) > maximum_pages:
        errors.append(f"ATS PDF exceeds maximum page count: {maximum_pages}")
    for index, page in enumerate(reader.pages, start=1):
        width = float(page.mediabox.width)
        height = float(page.mediabox.height)
        if abs(width - PAGE_WIDTH) > 0.1 or abs(height - PAGE_HEIGHT) > 0.1:
            errors.append(f"ATS PDF page {index} is not A4: {width:.2f} x {height:.2f} points")
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    errors.extend(_ats_text_errors(text, "ATS PDF", required_keywords))
    return errors


def _ats_text_errors(text: str, label: str, required_keywords: tuple[str, ...]) -> list[str]:
    errors: list[str] = []
    required = (
        "Profil professionnel",
        "Compétences",
        "Expérience professionnelle",
        "comptabilité",
        "150 %",
        *required_keywords,
    )
    for value in required:
        if value.casefold() not in text.casefold():
            errors.append(f"{label} is missing required text: {value}")
    positions = [text.casefold().find(value.casefold()) for value in ("Profil professionnel", "Compétences", "Expérience professionnelle", "Formation", "Langues")]
    if any(position < 0 for position in positions) or positions != sorted(positions):
        errors.append(f"{label} sections are not extracted in the expected ATS order")
    for forbidden in FORBIDDEN_CV_TEXT:
        if forbidden.casefold() in text.casefold():
            errors.append(f"{label} contains forbidden text: {forbidden}")
    if any(ord(character) < 32 and character not in "\n\r\t" for character in text):
        errors.append(f"{label} extraction contains control characters")
    return errors


def _check_ats_docx(path: Path, required_keywords: tuple[str, ...]) -> list[str]:
    errors: list[str] = []
    document = Document(str(path))
    errors.extend(_ats_text_errors(_document_text(document), "ATS DOCX", required_keywords))
    if document.tables:
        errors.append("ATS DOCX must not contain tables")
    if document.inline_shapes:
        errors.append("ATS DOCX must not contain images")
    if len(document.sections) != 1:
        errors.append("ATS DOCX must use a single document section")
    return errors


def check_documents(root: Path, public_dir: Path | None = None) -> dict[str, Any]:
    root = root.resolve()
    public_dir = public_dir or root / "site" / "assets" / "downloads"
    if not public_dir.is_absolute():
        public_dir = root / public_dir
    published = public_dir / STANDARD_CV_FILENAME
    template, strategy = _inputs(root)
    view = compose_standard_cv(root, template, strategy)
    state = DocumentCheckState(root=root, published=published, template=template, view=view)
    temporary = Path(tempfile.gettempdir()) / "delia-rossignol-life" / f"document-check-{uuid.uuid4().hex}"
    temporary.mkdir(parents=True)
    ats_results: list[dict[str, Any]] = []
    try:
        first = _generate_check_copies(state, temporary)
        reader = PdfReader(str(first))
        _check_pdf_pages(state, reader)
        _check_pdf_text(state, reader)
        _check_rendered_layout(state)
        ats_results, ats_errors = _check_ats_reproducibility(root, public_dir, temporary)
        state.errors.extend(ats_errors)
        for variant in ATS_VARIANTS:
            state.errors.extend(
                _check_ats_pdf(root, temporary / "ats-first" / variant.pdf_filename, variant.required_keywords)
            )
            state.errors.extend(
                _check_ats_docx(temporary / "ats-first" / variant.docx_filename, variant.required_keywords)
            )
    finally:
        shutil.rmtree(temporary, ignore_errors=True)
    report = state.report()
    report["documents"].extend(ats_results)
    report["errors"] = state.errors
    report["ok"] = not state.errors
    return report
