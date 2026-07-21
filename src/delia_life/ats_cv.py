from __future__ import annotations

import copy
import os
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import getDescent
from reportlab.pdfgen import canvas

from .core import load_json, sha256_file
from .cv_model import CVEducation, EvidenceRef
from .errors import ValidationError
from .pdf_renderer import line_wrap

ATS_PDF_FILENAME = "cv-delia-rossignol-ats.pdf"
ATS_DOCX_FILENAME = "cv-delia-rossignol-ats.docx"
ATS_TEMPLATE_PATH = Path("templates/cv/ats-classic/template.json")
ATS_STRATEGY_PATH = Path("data/style/cv-ats.json")
ATS_VARIANTS_PATH = Path("data/style/cv-ats-variants.json")
SKILLS_PATH = Path("data/knowledge/skills.json")

PAGE_WIDTH, PAGE_HEIGHT = A4
LEFT = 42.5
RIGHT = PAGE_WIDTH - 42.5
TOP = PAGE_HEIGHT - 42.5
BOTTOM = 42.5
CONTENT_WIDTH = RIGHT - LEFT
INK = HexColor("#162a3e")
NAVY = HexColor("#0d2f5a")
MUTED = HexColor("#5f6872")


@dataclass(frozen=True, slots=True)
class ATSCVVariant:
    id: str
    label: str
    description: str
    filename_stem: str
    required_keywords: tuple[str, ...]

    @property
    def pdf_filename(self) -> str:
        return f"{self.filename_stem}.pdf"

    @property
    def docx_filename(self) -> str:
        return f"{self.filename_stem}.docx"


ATS_VARIANTS = (
    ATSCVVariant(
        id="transverse",
        label="ATS transverse",
        description="Conseil, commerce, gestion administrative et coordination réunis dans un même CV.",
        filename_stem="cv-delia-rossignol-ats",
        required_keywords=("Conseil", "relation client", "Commerce", "Gestion administrative", "Coordination de projets"),
    ),
    ATSCVVariant(
        id="relation-client",
        label="Conseil & relation client",
        description="Accompagnement, expérience client, réclamations et fidélisation.",
        filename_stem="cv-delia-rossignol-ats-relation-client",
        required_keywords=("Conseil", "relation client", "fidélisation", "réclamations", "clientèle"),
    ),
    ATSCVVariant(
        id="commerce-vente",
        label="Commerce & vente",
        description="Vente-conseil, développement commercial, négociation et performance.",
        filename_stem="cv-delia-rossignol-ats-commerce-vente",
        required_keywords=("Commerce", "vente", "fidélisation", "80 %", "négocier"),
    ),
    ATSCVVariant(
        id="gestion-administrative",
        label="Gestion administrative",
        description="Suivi d’activité, facturation, fournisseurs et coordination de projets.",
        filename_stem="cv-delia-rossignol-ats-gestion-administrative",
        required_keywords=("Gestion administrative", "facturation", "Coordination de projets", "fournisseurs", "contrats"),
    ),
    ATSCVVariant(
        id="luxe-mode-relation-client",
        label="Luxe, mode & relation client",
        description="Vente-conseil, accompagnement personnalisé et clientèles nationales et internationales.",
        filename_stem="cv-delia-rossignol-ats-luxe-mode-relation-client",
        required_keywords=("luxe", "mode", "prêt-à-porter", "accompagnement personnalisé", "clientèle"),
    ),
)


def ats_variant(identifier: str) -> ATSCVVariant:
    try:
        return next(variant for variant in ATS_VARIANTS if variant.id == identifier)
    except StopIteration as error:
        raise ValidationError(f"Unknown ATS CV variant: {identifier}") from error


def load_ats_strategy(root: Path, variant_id: str) -> dict[str, Any]:
    strategy = load_json(root / ATS_STRATEGY_PATH)
    if variant_id == "transverse":
        return strategy
    catalog = load_json(root / ATS_VARIANTS_PATH)
    if catalog.get("status") != "validated":
        raise ValidationError("ATS CV variants must be validated")
    overlays = {str(item["id"]): item for item in catalog["variants"]}
    if set(overlays) != {variant.id for variant in ATS_VARIANTS if variant.id != "transverse"}:
        raise ValidationError("ATS CV variant catalog does not match the generator catalog")
    try:
        overlay = overlays[variant_id]
    except KeyError as error:
        raise ValidationError(f"ATS CV variant content is missing: {variant_id}") from error
    result = copy.deepcopy(strategy)
    for key in ("strategy_id", "target_title", "profile_summary", "skill_ids"):
        target_key = "id" if key == "strategy_id" else key
        result[target_key] = copy.deepcopy(overlay[key])
    if "profile_summary_evidence" in overlay:
        result["profile_summary_evidence"] = copy.deepcopy(overlay["profile_summary_evidence"])
    experiences = {str(item["entity_id"]): item for item in result["experiences"]}
    for override in overlay["experience_overrides"]:
        entity_id = str(override["entity_id"])
        try:
            experience = experiences[entity_id]
        except KeyError as error:
            raise ValidationError(f"ATS CV variant overrides an unknown experience: {entity_id}") from error
        experience["bullets"] = copy.deepcopy(override["bullets"])
    return result


@dataclass(frozen=True, slots=True)
class ATSBullet:
    text: str
    evidence: tuple[EvidenceRef, ...]


@dataclass(frozen=True, slots=True)
class ATSExperience:
    entity_id: str
    heading: str
    period: str
    mission: str
    bullets: tuple[ATSBullet, ...]
    evidence: tuple[EvidenceRef, ...]


@dataclass(frozen=True, slots=True)
class ATSCVViewModel:
    template_id: str
    template_version: str
    strategy_id: str
    name: str
    email: str
    phone: str
    target_title: str
    profile_summary: str
    skills: tuple[str, ...]
    experiences: tuple[ATSExperience, ...]
    education: tuple[CVEducation, ...]
    language: str
    evidence: tuple[EvidenceRef, ...]

    @property
    def source_ids(self) -> tuple[str, ...]:
        return tuple(sorted({item.source_id for item in self.evidence}))


@dataclass
class KnowledgeResolver:
    root: Path
    cache: dict[tuple[str, str], dict[str, Any]] = field(default_factory=dict)

    def entity(self, entity_type: str, entity_id: str) -> dict[str, Any]:
        key = (entity_type, entity_id)
        if key in self.cache:
            return self.cache[key]
        candidates = (
            self.root / "data" / "knowledge" / "entities" / entity_type / f"{entity_id}.json",
            self.root / "data" / "knowledge" / "reference" / entity_type / f"{entity_id}.json",
        )
        for path in candidates:
            if path.is_file():
                document = load_json(path)
                self.cache[key] = document
                return document
        raise ValidationError(f"Knowledge entity is missing: {entity_type}/{entity_id}")

    def field(self, entity_type: str, entity_id: str, field_name: str) -> tuple[Any, tuple[EvidenceRef, ...]]:
        entity = self.entity(entity_type, entity_id)
        envelope = entity.get("fields", {}).get(field_name)
        if not isinstance(envelope, dict) or "value" not in envelope:
            raise ValidationError(f"Validated field is missing: {entity_type}/{entity_id}/{field_name}")
        evidence = tuple(
            EvidenceRef(
                proposal_id=str(item["proposal_id"]),
                source_id=str(item["source_id"]),
                locator=str(item["locator"]),
            )
            for item in envelope.get("provenance", [])
        )
        if not evidence:
            raise ValidationError(f"Validated field has no provenance: {entity_type}/{entity_id}/{field_name}")
        return envelope["value"], evidence

    def referenced(self, specification: dict[str, Any]) -> tuple[Any, tuple[EvidenceRef, ...]]:
        return self.field(
            str(specification["entity_type"]),
            str(specification["entity_id"]),
            str(specification["field"]),
        )

    def evidence_for(self, specifications: list[dict[str, Any]]) -> tuple[EvidenceRef, ...]:
        return tuple(reference for specification in specifications for reference in self.referenced(specification)[1])

    def experience_period(self, entity_id: str) -> tuple[str, tuple[EvidenceRef, ...]]:
        entity = self.entity("experience", entity_id)
        for field_name in ("timeframe", "chronology", "details"):
            if field_name not in entity.get("fields", {}):
                continue
            value, evidence = self.field("experience", entity_id, field_name)
            if not isinstance(value, dict):
                continue
            start = value.get("start_date", value.get("start_year"))
            end = value.get("end_date", value.get("end_year", value.get("known_through")))
            if start is not None and end is not None:
                return f"{_date(start)} - {_date(end)}", evidence
        raise ValidationError(f"Experience period cannot be derived: {entity_id}")

    def education(self, entity_type: str, entity_id: str) -> CVEducation:
        details, evidence = self.field(entity_type, entity_id, "details")
        if not isinstance(details, dict):
            raise ValidationError(f"Education details must be structured: {entity_id}")
        if "awarded_year" in details:
            primary = f"{details['awarded_year']} - {details['name']} - {details['level']}"
            secondary = str(details["issuer"])
        elif "program" in details:
            primary = f"{str(details.get('end_date', details.get('start_date', '')))[:4]} - {details['program']}"
            secondary = str(details["institution"])
        elif "program_stage" in details:
            primary = f"{details['year']} - {str(details['program_stage']).capitalize()}"
            secondary = str(details["institution"])
        else:
            raise ValidationError(f"Unsupported education structure: {entity_id}")
        return CVEducation(entity_id=entity_id, primary=primary, secondary=secondary, evidence=evidence)


def _date(value: Any) -> str:
    raw = str(value)
    if len(raw) >= 7 and raw[4] == "-":
        return f"{raw[5:7]}/{raw[:4]}"
    return raw[:4]


def _profile(resolver: KnowledgeResolver, strategy: dict[str, Any]) -> tuple[str, str, str, tuple[EvidenceRef, ...]]:
    name, name_evidence = resolver.field("person", str(strategy["person_id"]), "professional_name")
    email, email_evidence = resolver.field("contact-point", str(strategy["email_contact_id"]), "details")
    phone, phone_evidence = resolver.field("contact-point", str(strategy["phone_contact_id"]), "details")
    return (
        str(name),
        str(email["value"]),
        str(phone["value"]),
        tuple((*name_evidence, *email_evidence, *phone_evidence)),
    )


def _skills(
    root: Path,
    resolver: KnowledgeResolver,
    strategy: dict[str, Any],
    maximum: int,
) -> tuple[tuple[str, ...], tuple[EvidenceRef, ...]]:
    index = load_json(root / SKILLS_PATH)
    if index.get("status") != "validated":
        raise ValidationError("Transferable skills index must be validated")
    by_id = {str(item["id"]): item for item in index["skills"]}
    selected: list[str] = []
    evidence: list[EvidenceRef] = []
    for skill_id in strategy["skill_ids"]:
        try:
            skill = by_id[str(skill_id)]
        except KeyError as error:
            raise ValidationError(f"ATS CV references an unknown skill: {skill_id}") from error
        if not skill.get("transferable"):
            raise ValidationError(f"ATS CV skill is not transferable: {skill_id}")
        selected.append(str(skill["name"]))
        evidence.extend(resolver.evidence_for(skill["evidence"]))
    if len(selected) > maximum:
        raise ValidationError("ATS CV skills exceed template limit")
    return tuple(selected), tuple(evidence)


def _experiences(
    resolver: KnowledgeResolver,
    strategy: dict[str, Any],
    maximum_bullets: int,
    maximum_words: int,
) -> tuple[ATSExperience, ...]:
    experiences: list[ATSExperience] = []
    for specification in strategy["experiences"]:
        entity_id = str(specification["entity_id"])
        mission, mission_evidence = resolver.field("experience", entity_id, "mission")
        period, period_evidence = resolver.experience_period(entity_id)
        heading_evidence = resolver.evidence_for(specification["heading_evidence"])
        bullet_specs = specification["bullets"]
        if len(bullet_specs) > maximum_bullets:
            raise ValidationError(f"ATS experience exceeds bullet limit: {entity_id}")
        bullets: list[ATSBullet] = []
        for bullet_spec in bullet_specs:
            text = str(bullet_spec["text"])
            if len(text.split()) > maximum_words:
                raise ValidationError(f"ATS bullet exceeds word limit: {entity_id}")
            bullet_evidence = resolver.evidence_for(bullet_spec["evidence"])
            bullets.append(ATSBullet(text=text, evidence=bullet_evidence))
        all_evidence = tuple(
            (
                *mission_evidence,
                *period_evidence,
                *heading_evidence,
                *(reference for bullet in bullets for reference in bullet.evidence),
            )
        )
        experiences.append(
            ATSExperience(
                entity_id=entity_id,
                heading=str(specification["heading"]),
                period=period,
                mission=str(mission),
                bullets=tuple(bullets),
                evidence=all_evidence,
            )
        )
    return tuple(experiences)


def compose_ats_cv(root: Path, template: dict[str, Any], strategy: dict[str, Any]) -> ATSCVViewModel:
    if strategy.get("status") != "validated":
        raise ValidationError("ATS CV content strategy must be validated")
    if strategy.get("template_id") != template.get("id"):
        raise ValidationError("ATS CV content strategy and template do not match")
    if template.get("rendering", {}).get("engine") != "standard-single-column-v1":
        raise ValidationError("ATS CV requires the single-column rendering engine")
    rules = template["content_rules"]
    summary = str(strategy["profile_summary"])
    if len(summary.split()) > int(rules["profile_max_words"]):
        raise ValidationError("ATS CV profile exceeds template word limit")
    resolver = KnowledgeResolver(root.resolve())
    name, email, phone, profile_evidence = _profile(resolver, strategy)
    summary_evidence = resolver.evidence_for(strategy.get("profile_summary_evidence", []))
    skills, skill_evidence = _skills(root, resolver, strategy, int(rules["key_skills_max"]))
    experiences = _experiences(
        resolver,
        strategy,
        int(rules["experience_bullets_max"]),
        int(rules["bullet_max_words"]),
    )
    education = tuple(
        resolver.education(str(item["entity_type"]), str(item["entity_id"])) for item in strategy["education"]
    )
    language, language_evidence = resolver.field("language-proficiency", str(strategy["language_id"]), "details")
    language_line = (
        f"{str(language['language']).capitalize()} - oral et lecture : {language['speaking_level']} - "
        f"écrit : {language['writing_level']}"
    )
    evidence = tuple(
        (
            *profile_evidence,
            *summary_evidence,
            *skill_evidence,
            *language_evidence,
            *(reference for item in experiences for reference in item.evidence),
            *(reference for item in education for reference in item.evidence),
        )
    )
    return ATSCVViewModel(
        template_id=str(template["id"]),
        template_version=str(template["version"]),
        strategy_id=str(strategy["id"]),
        name=name,
        email=email,
        phone=phone,
        target_title=str(strategy["target_title"]),
        profile_summary=summary,
        skills=skills,
        experiences=experiences,
        education=education,
        language=language_line,
        evidence=evidence,
    )


@dataclass
class PDFFlow:
    pdf: canvas.Canvas
    y: float = TOP
    page: int = 1
    minimum_y: float = TOP

    def new_page(self) -> None:
        self.pdf.showPage()
        self.page += 1
        self.y = TOP

    def ensure(self, height: float) -> None:
        if self.y - height < BOTTOM:
            self.new_page()

    def text(
        self,
        value: str,
        *,
        font: str = "Helvetica",
        size: float = 9.2,
        leading: float = 11.6,
        color: Any = INK,
        left: float = LEFT,
        width: float = CONTENT_WIDTH,
        before: float = 0,
        after: float = 0,
    ) -> None:
        lines = line_wrap(value, font, size, width)
        height = before + (len(lines) * leading) + after
        self.ensure(height)
        self.y -= before
        self.pdf.setFont(font, size)
        self.pdf.setFillColor(color)
        for line in lines:
            self.pdf.drawString(left, self.y, line)
            self.y -= leading
        self.y -= after
        if lines:
            descent = abs(getDescent(font, size))
            self.minimum_y = min(self.minimum_y, self.y + after + leading - descent)


def _section(flow: PDFFlow, title: str) -> None:
    flow.ensure(28)
    flow.text(title.upper(), font="Helvetica-Bold", size=10.2, leading=12, color=NAVY, before=8, after=3)
    flow.pdf.setStrokeColor(NAVY)
    flow.pdf.setLineWidth(0.6)
    flow.pdf.line(LEFT, flow.y + 8, RIGHT, flow.y + 8)
    flow.y -= 5


def _estimated_experience_height(experience: ATSExperience) -> float:
    heading_lines = len(line_wrap(f"{experience.heading} | {experience.period}", "Helvetica-Bold", 9.3, CONTENT_WIDTH))
    mission_lines = len(line_wrap(experience.mission, "Helvetica-Oblique", 8.8, CONTENT_WIDTH))
    bullet_lines = sum(len(line_wrap(f"- {bullet.text}", "Helvetica", 8.8, CONTENT_WIDTH - 10)) for bullet in experience.bullets)
    return (heading_lines * 11.2) + (mission_lines * 10.8) + (bullet_lines * 10.8) + 15


def render_ats_pdf(view: ATSCVViewModel, output: Path) -> dict[str, Any]:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.name}.{uuid.uuid4().hex}.tmp")
    try:
        pdf = canvas.Canvas(str(temporary), pagesize=A4, pageCompression=1, invariant=1)
        pdf.setTitle("CV Délia Rossignol - ATS")
        pdf.setAuthor("Délia Rossignol")
        flow = PDFFlow(pdf)
        flow.text(view.name.upper(), font="Helvetica-Bold", size=17, leading=20, color=NAVY, after=2)
        flow.text(view.target_title, font="Helvetica-Bold", size=10.5, leading=13, color=INK, after=2)
        flow.text(f"{view.email} | {view.phone}", size=9, leading=11, color=MUTED, after=5)

        _section(flow, "Profil professionnel")
        flow.text(view.profile_summary, size=9.2, leading=11.8, after=2)

        _section(flow, "Compétences")
        flow.text("; ".join(view.skills), size=9, leading=11.5, after=1)

        _section(flow, "Expérience professionnelle")
        for experience in view.experiences:
            flow.ensure(_estimated_experience_height(experience))
            flow.text(
                f"{experience.heading} | {experience.period}",
                font="Helvetica-Bold",
                size=9.3,
                leading=11.2,
                color=NAVY,
                before=4,
                after=1,
            )
            flow.text(experience.mission, font="Helvetica-Oblique", size=8.8, leading=10.8, after=1)
            for bullet in experience.bullets:
                flow.text(
                    f"- {bullet.text}",
                    size=8.8,
                    leading=10.8,
                    left=LEFT + 10,
                    width=CONTENT_WIDTH - 10,
                    after=1,
                )

        _section(flow, "Formation")
        for education in view.education:
            flow.text(f"{education.primary} | {education.secondary}", size=8.9, leading=11, after=1)

        _section(flow, "Langues")
        flow.text(view.language, size=8.9, leading=11)
        pages = flow.page
        pdf.save()
        if flow.minimum_y < BOTTOM:
            raise ValidationError("ATS PDF content crosses the safe lower page boundary")
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)
    return {"pages": pages, "minimum_y": flow.minimum_y, "overflow_count": 0}


def _set_docx_defaults(document: Any) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor(22, 42, 62)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.widow_control = True

    heading_one = document.styles["Heading 1"]
    heading_one.font.name = "Arial"
    heading_one.font.size = Pt(10.5)
    heading_one.font.bold = True
    heading_one.font.color.rgb = RGBColor(13, 47, 90)
    heading_one.paragraph_format.space_before = Pt(7)
    heading_one.paragraph_format.space_after = Pt(3)
    heading_one.paragraph_format.keep_with_next = True

    heading_two = document.styles["Heading 2"]
    heading_two.font.name = "Arial"
    heading_two.font.size = Pt(9.5)
    heading_two.font.bold = True
    heading_two.font.color.rgb = RGBColor(13, 47, 90)
    heading_two.paragraph_format.space_before = Pt(4)
    heading_two.paragraph_format.space_after = Pt(1)
    heading_two.paragraph_format.keep_with_next = True

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(9.5)
    bullet.paragraph_format.left_indent = Cm(0.5)
    bullet.paragraph_format.first_line_indent = Cm(-0.25)
    bullet.paragraph_format.space_after = Pt(1)
    bullet.paragraph_format.line_spacing = 1.0


def _set_run_font(run: Any, *, size: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _docx_heading(document: Any, title: str) -> None:
    document.add_paragraph(title.upper(), style="Heading 1")


def _normalize_docx(source: Path, destination: Path) -> None:
    fixed_timestamp = (1980, 1, 1, 0, 0, 0)
    temporary = destination.with_name(f".{destination.name}.{uuid.uuid4().hex}.tmp")
    try:
        with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
            temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
        ) as normalized:
            for name in sorted(archive.namelist()):
                info = zipfile.ZipInfo(name, fixed_timestamp)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                normalized.writestr(info, archive.read(name))
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)


def render_ats_docx(view: ATSCVViewModel, output: Path) -> dict[str, Any]:
    output.parent.mkdir(parents=True, exist_ok=True)
    raw = output.with_name(f".{output.name}.{uuid.uuid4().hex}.raw.docx")
    try:
        document = Document()
        _set_docx_defaults(document)
        fixed_date = datetime(2026, 7, 20, tzinfo=UTC)
        document.core_properties.author = "Délia Rossignol"
        document.core_properties.title = "CV Délia Rossignol - ATS"
        document.core_properties.created = fixed_date
        document.core_properties.modified = fixed_date

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(view.name.upper())
        _set_run_font(run, size=17, bold=True, color=RGBColor(13, 47, 90))

        for value, bold in ((view.target_title, True), (f"{view.email} | {view.phone}", False)):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(value)
            _set_run_font(run, size=10 if bold else 9, bold=bold)

        _docx_heading(document, "Profil professionnel")
        document.add_paragraph(view.profile_summary)
        _docx_heading(document, "Compétences")
        document.add_paragraph("; ".join(view.skills))
        _docx_heading(document, "Expérience professionnelle")
        for experience in view.experiences:
            paragraph = document.add_paragraph(style="Heading 2")
            run = paragraph.add_run(f"{experience.heading} | {experience.period}")
            _set_run_font(run, size=9.5, bold=True, color=RGBColor(13, 47, 90))
            mission = document.add_paragraph()
            mission.paragraph_format.space_after = Pt(1)
            mission.paragraph_format.keep_with_next = True
            mission_run = mission.add_run(experience.mission)
            mission_run.italic = True
            for bullet in experience.bullets:
                document.add_paragraph(bullet.text, style="List Bullet")

        _docx_heading(document, "Formation")
        for education in view.education:
            document.add_paragraph(f"{education.primary} | {education.secondary}")
        _docx_heading(document, "Langues")
        document.add_paragraph(view.language)
        document.save(str(raw))
        _normalize_docx(raw, output)
    finally:
        raw.unlink(missing_ok=True)
    return {"tables": 0, "images": 0}


def build_ats_cv(
    root: Path,
    pdf_output: Path,
    docx_output: Path,
    *,
    variant_id: str = "transverse",
) -> list[dict[str, Any]]:
    root = root.resolve()
    variant = ats_variant(variant_id)
    template = load_json(root / ATS_TEMPLATE_PATH)
    strategy = load_ats_strategy(root, variant_id)
    view = compose_ats_cv(root, template, strategy)
    pdf_layout = render_ats_pdf(view, pdf_output)
    docx_layout = render_ats_docx(view, docx_output)
    common = {
        "template": view.template_id,
        "template_version": view.template_version,
        "content_strategy": view.strategy_id,
        "source_ids": list(view.source_ids),
        "variant": variant.id,
        "variant_label": variant.label,
    }
    document_id = "standard-cv-ats" if variant.id == "transverse" else f"standard-cv-ats-{variant.id}"
    return [
        {
            **common,
            "id": f"{document_id}-pdf",
            "format": "pdf",
            "output": str(pdf_output.resolve()),
            "pages": int(pdf_layout["pages"]),
            "sha256": sha256_file(pdf_output),
            "layout": pdf_layout,
        },
        {
            **common,
            "id": f"{document_id}-docx",
            "format": "docx",
            "output": str(docx_output.resolve()),
            "sha256": sha256_file(docx_output),
            "layout": docx_layout,
        },
    ]


def build_all_ats_cvs(root: Path, output_dir: Path) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for variant in ATS_VARIANTS:
        results.extend(
            build_ats_cv(
                root,
                output_dir / variant.pdf_filename,
                output_dir / variant.docx_filename,
                variant_id=variant.id,
            )
        )
    return results
