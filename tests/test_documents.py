from __future__ import annotations

import sys
import unittest
import uuid
from pathlib import Path
from tempfile import gettempdir
from unittest.mock import patch

from docx import Document
from PIL import Image
from pypdf import PdfReader
from pypdf.generic import ContentStream
from reportlab.pdfbase.pdfmetrics import getAscent, stringWidth

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
TEST_TMP = Path(gettempdir()) / "delia-rossignol-life-tests"
TEST_TMP.mkdir(exist_ok=True)

from delia_life.ats_cv import (
    ATS_DOCX_FILENAME,
    ATS_PDF_FILENAME,
    ATS_VARIANTS,
    build_all_ats_cvs,
    build_ats_cv,
    compose_ats_cv,
    load_ats_strategy,
)
from delia_life.core import load_json
from delia_life.cv_composer import compose_standard_cv
from delia_life.document_builder import build_documents, build_standard_cv, check_documents
from delia_life.errors import ValidationError
from delia_life.pdf_layout import CVLayoutRules, calculate_card_geometry
from delia_life.pdf_renderer import LEFT, RIGHT, composite_photo_background, line_wrap, render_standard_cv
from delia_life.storage import remove_tree


class DocumentTests(unittest.TestCase):
    def setUp(self) -> None:
        self.work = TEST_TMP / f"{self._testMethodName}-{uuid.uuid4().hex}"
        self.work.mkdir(parents=True)

    def tearDown(self) -> None:
        if self.work.exists():
            remove_tree(self.work, ignore_errors=True)

    def test_standard_cv_is_reproducible_a4_and_complete(self) -> None:
        first = self.work / "first.pdf"
        second = self.work / "second.pdf"
        first_result = build_standard_cv(ROOT, first)
        second_result = build_standard_cv(ROOT, second)
        self.assertEqual(first.read_bytes(), second.read_bytes())
        self.assertEqual(first_result["sha256"], second_result["sha256"])

        reader = PdfReader(str(first))
        self.assertEqual(len(reader.pages), 2)
        for page in reader.pages:
            self.assertAlmostEqual(float(page.mediabox.width), 595.2756, places=2)
            self.assertAlmostEqual(float(page.mediabox.height), 841.8898, places=2)
        page_texts = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(page_texts)
        self.assertIn("deliarossignol@gmail.com", text)
        self.assertIn("06 20 67 40 52", text)
        self.assertIn("construire ensemble une solution concrète", text)
        for sequence_label in ["CADRER", "CONCEVOIR", "VENDRE & ENGAGER", "PILOTER & LIVRER"]:
            self.assertIn(sequence_label, text)
        self.assertIn("CONCEPTION & PILOTAGE DE PROJETS D’AGENCEMENT", text)
        self.assertIn("Même métier, deux contextes successifs", text)
        self.assertIn("SOCLE MÉTIER COMMUN", text)
        self.assertIn("BLEU ROSSIGNOL", text)
        self.assertIn("PERFECTEUR D’INTÉRIEUR", text)
        self.assertIn("acquisition & fidélisation", text)
        self.assertIn("Analytics · 3 optimisations SEO", text)
        self.assertIn("sourcing · évaluation ·", text)
        self.assertIn("négociation fournisseurs", text)
        self.assertNotIn("CV standard", text)
        self.assertNotIn("Signature éditoriale", text)
        self.assertIn("1 / 2", page_texts[0])
        self.assertIn("2 / 2", page_texts[1])
        self.assertEqual(page_texts[1].count("PÉRIMÈTRE"), 6)
        for complementary_detail in (
            "Gérer 3 mandats immobiliers",
            "relations bancaires",
            "100 couverts par jour",
            "objectif commercial",
            "Flowerbomb et Amor Amor",
            "avenue Montaigne",
        ):
            self.assertIn(complementary_detail, page_texts[1])
        self.assertNotIn("41 ans", text)
        self.assertNotIn("2 enfants", text)

    def test_published_standard_cv_is_current(self) -> None:
        result = check_documents(ROOT)
        self.assertTrue(result["ok"], result["errors"])

    def test_document_build_publishes_the_same_atomic_artifact(self) -> None:
        result = build_documents(ROOT, self.work / "output", self.work / "public")
        self.assertEqual(
            [document["id"] for document in result["documents"]],
            [
                "standard-cv-signature-editorial",
                "standard-cv-ats-pdf",
                "standard-cv-ats-docx",
                "standard-cv-ats-relation-client-pdf",
                "standard-cv-ats-relation-client-docx",
                "standard-cv-ats-commerce-vente-pdf",
                "standard-cv-ats-commerce-vente-docx",
                "standard-cv-ats-gestion-administrative-pdf",
                "standard-cv-ats-gestion-administrative-docx",
            ],
        )
        for document in result["documents"]:
            output = Path(document["output"])
            public = Path(document["public_output"])
            self.assertEqual(output.read_bytes(), public.read_bytes())
            self.assertEqual(document["sha256"], __import__("hashlib").sha256(output.read_bytes()).hexdigest())

    def test_ats_cv_is_reproducible_traceable_and_extracted_in_section_order(self) -> None:
        first_dir = self.work / "first"
        second_dir = self.work / "second"
        first = build_ats_cv(ROOT, first_dir / ATS_PDF_FILENAME, first_dir / ATS_DOCX_FILENAME)
        build_ats_cv(ROOT, second_dir / ATS_PDF_FILENAME, second_dir / ATS_DOCX_FILENAME)
        self.assertEqual(
            (first_dir / ATS_PDF_FILENAME).read_bytes(),
            (second_dir / ATS_PDF_FILENAME).read_bytes(),
        )
        self.assertEqual(
            (first_dir / ATS_DOCX_FILENAME).read_bytes(),
            (second_dir / ATS_DOCX_FILENAME).read_bytes(),
        )
        self.assertTrue(all(document["source_ids"] for document in first))

        reader = PdfReader(first_dir / ATS_PDF_FILENAME)
        self.assertEqual(len(reader.pages), 1)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        sections = [
            "PROFIL PROFESSIONNEL",
            "COMP\u00c9TENCES",
            "EXP\u00c9RIENCE PROFESSIONNELLE",
            "FORMATION",
            "LANGUES",
        ]
        positions = [text.casefold().find(section.casefold()) for section in sections]
        self.assertTrue(all(position >= 0 for position in positions), (positions, text))
        self.assertEqual(positions, sorted(positions))
        self.assertIn("Google Analytics", text)
        self.assertIn("comptabilit\u00e9", text)
        self.assertIn("150 %", text)
        self.assertFalse(any(ord(character) < 32 and character not in "\n\r\t" for character in text))

        document = Document(first_dir / ATS_DOCX_FILENAME)
        docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Gestion administrative", docx_text)
        self.assertFalse(document.tables)
        self.assertFalse(document.inline_shapes)
        self.assertGreaterEqual(sum(paragraph.style.name == "List Bullet" for paragraph in document.paragraphs), 10)

    def test_ats_cv_uses_only_validated_transferable_skills(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "ats-classic" / "template.json")
        strategy = load_json(ROOT / "data" / "style" / "cv-ats.json")
        view = compose_ats_cv(ROOT, template, strategy)
        self.assertEqual(view.skills[0], "Relation et service client")
        self.assertEqual(len(view.skills), 8)
        self.assertEqual(view.experiences[0].entity_id, "bleu-rossignol-founder")
        self.assertGreater(len(view.source_ids), 5)

        strategy["skill_ids"][0] = "unknown-skill"
        with self.assertRaisesRegex(ValueError, "unknown skill"):
            compose_ats_cv(ROOT, template, strategy)

    def test_all_ats_variants_are_one_page_and_use_distinct_positioning(self) -> None:
        results = build_all_ats_cvs(ROOT, self.work)
        self.assertEqual(len(results), len(ATS_VARIANTS) * 2)
        template = load_json(ROOT / "templates" / "cv" / "ats-classic" / "template.json")
        titles: set[str] = set()
        expected_first_skills = {
            "transverse": "Relation et service client",
            "relation-client": "Relation et service client",
            "commerce-vente": "Vente-conseil et négociation",
            "gestion-administrative": "Gestion administrative",
        }
        for variant in ATS_VARIANTS:
            strategy = load_ats_strategy(ROOT, variant.id)
            view = compose_ats_cv(ROOT, template, strategy)
            titles.add(view.target_title)
            self.assertEqual(view.skills[0], expected_first_skills[variant.id])
            reader = PdfReader(self.work / variant.pdf_filename)
            self.assertEqual(len(reader.pages), 1, variant.id)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            for keyword in variant.required_keywords:
                self.assertIn(keyword.casefold(), text.casefold(), (variant.id, keyword))
        self.assertEqual(len(titles), len(ATS_VARIANTS))

    def test_cv_composition_is_traceable_and_derives_periods(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        strategy = load_json(ROOT / "data" / "style" / "cv-standard.json")
        view = compose_standard_cv(ROOT, template, strategy)
        self.assertEqual(
            view.tagline,
            "Conseil · Commerce · Gestion administrative · Gestion de projet",
        )
        self.assertGreater(len(view.source_ids), 5)
        group = view.recent_continuity_groups[0]
        self.assertEqual(group.period, "08/2020 - 07/2026")
        self.assertEqual(
            tuple(item.entity_id for item in group.experiences),
            ("bleu-rossignol-founder", "raison-home-franchisee"),
        )
        self.assertEqual(group.experiences[0].organization_name, "BLEU ROSSIGNOL")
        self.assertEqual(group.experiences[0].organization_tagline, "PERFECTEUR D’INTÉRIEUR")
        self.assertEqual(group.experiences[1].organization_name, "RAISON HOME")
        highlight = group.experiences[0].highlight
        self.assertIsNotNone(highlight)
        assert highlight is not None
        self.assertEqual(
            tuple(card.label for card in highlight.cards),
            ("Développer l’activité", "Piloter le digital", "Structurer l’écosystème"),
        )
        self.assertEqual(tuple(len(card.items) for card in highlight.cards), (2, 2, 2))
        sequence = group.responsibility_sequence
        self.assertEqual(
            tuple(step.label for step in sequence),
            ("Cadrer", "Concevoir", "Vendre & engager", "Piloter & livrer"),
        )
        self.assertEqual(
            tuple(responsibility for step in sequence for responsibility in step.responsibilities),
            (
                "analyse du besoin client",
                "relevé de mesures et prise de cotes",
                "conception fonctionnelle",
                "choix esthétique",
                "rendus 3D",
                "plans techniques d’exécution",
                "chiffrage et devis",
                "négociation commerciale",
                "commande fabricants",
                "planification",
                "coordination des artisans",
                "suivi de chantier",
                "contrôle de la pose",
                "réception du chantier",
                "gestion des réclamations",
                "facturation",
            ),
        )
        self.assertEqual(view.complementary_experiences[3].period, "11/2006 - 09/2008")
        self.assertEqual(
            tuple(item.bullet_limit for item in view.complementary_experiences),
            (2, 5, 2, 2, 2, 1),
        )
        self.assertTrue(all(experience.evidence for experience in (*view.recent_experiences, *view.complementary_experiences)))

    def test_cv_template_rules_are_enforced_by_the_composer(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        strategy = load_json(ROOT / "data" / "style" / "cv-standard.json")
        strategy["profile_summary"] = "mot " * (template["content_rules"]["profile_max_words"] + 1)
        with self.assertRaisesRegex(ValueError, "profile exceeds"):
            compose_standard_cv(ROOT, template, strategy)

    def test_cv_sequence_must_preserve_the_complete_responsibility_sequence(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        strategy = load_json(ROOT / "data" / "style" / "cv-standard.json")
        strategy["continuity_groups"][0]["responsibility_sequence"][-1]["responsibility_count"] = 6
        with self.assertRaisesRegex(ValueError, "cover every responsibility"):
            compose_standard_cv(ROOT, template, strategy)

    def test_cv_continuity_group_requires_a_validated_shared_job_role(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        strategy = load_json(ROOT / "data" / "style" / "cv-standard.json")
        strategy["continuity_groups"][0]["job_role_id"] = "unvalidated-role"
        with self.assertRaisesRegex(ValueError, "job role is not validated"):
            compose_standard_cv(ROOT, template, strategy)

    def test_pdf_renderer_contains_no_delia_business_content(self) -> None:
        source = (ROOT / "src" / "delia_life" / "pdf_renderer.py").read_text(encoding="utf-8")
        for personal_value in ["BLEU ROSSIGNOL", "Raison Home", "Cuisinella", "Winner", "deliarossignol@gmail.com"]:
            self.assertNotIn(personal_value, source)

    def test_cv_layout_rules_require_a_consistent_spacing_grid(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        values = dict(template["rendering"]["layout"])
        rules = CVLayoutRules.from_mapping(values)
        self.assertEqual(rules.component_gap_pt, rules.spacing_unit_pt * 2)
        self.assertEqual(rules.compact_experience_gap_pt, rules.spacing_unit_pt * 3)
        values["component_gap_pt"] = 7
        with self.assertRaisesRegex(ValueError, "multiple of spacing_unit_pt"):
            CVLayoutRules.from_mapping(values)

    def test_card_geometry_keeps_the_visible_stroke_inside_its_allocation(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        rules = CVLayoutRules.from_mapping(template["rendering"]["layout"])
        geometry = calculate_card_geometry(
            x=65,
            top=400,
            width=RIGHT - 65,
            line_count=2,
            label_size=7,
            label_leading=8,
            body_size=6.9,
            body_leading=8,
            rules=rules,
            minimum_height=40,
        )
        visible_left = geometry.path_x - (geometry.stroke_width / 2)
        visible_right = geometry.path_x + geometry.path_width + (geometry.stroke_width / 2)
        self.assertAlmostEqual(visible_left, geometry.outer_x)
        self.assertAlmostEqual(visible_right, geometry.outer_right)
        self.assertLessEqual(visible_right, RIGHT)
        self.assertEqual(geometry.label_to_divider_gap, geometry.divider_to_body_gap)

    def test_line_wrap_splits_a_single_word_that_exceeds_the_available_width(self) -> None:
        word = "SUPERCALIFRAGILISTICEXPIALIDOCIOUS"
        width = 50
        lines = line_wrap(word, "Helvetica", 8, width)
        self.assertEqual("".join(lines), word)
        self.assertTrue(all(stringWidth(line, "Helvetica", 8) <= width for line in lines))

    def test_failed_layout_audit_does_not_replace_an_existing_pdf(self) -> None:
        template = load_json(ROOT / "templates" / "cv" / "signature-editorial" / "template.json")
        strategy = load_json(ROOT / "data" / "style" / "cv-standard.json")
        view = compose_standard_cv(ROOT, template, strategy)
        output = self.work / "atomic.pdf"
        output.write_bytes(b"existing-valid-document")
        with (
            patch("delia_life.pdf_renderer.LayoutAudit.report", return_value={"violations": ["forced failure"]}),
            self.assertRaisesRegex(ValidationError, "forced failure"),
        ):
            render_standard_cv(view, output)
        self.assertEqual(output.read_bytes(), b"existing-valid-document")
        self.assertEqual(list(self.work.glob(".*.tmp")), [])

    def test_transparent_photo_is_composited_on_the_warm_cv_background(self) -> None:
        transparent = Image.new("RGBA", (10, 10), (0, 0, 0, 0))
        transparent.putpixel((5, 5), (25, 50, 75, 255))
        result = composite_photo_background(transparent)
        self.assertEqual(result.mode, "RGB")
        self.assertEqual(result.getpixel((0, 0)), (249, 245, 243))
        self.assertEqual(result.getpixel((5, 5)), (25, 50, 75))

    def test_generated_cv_layout_audit_enforces_equal_component_gaps(self) -> None:
        result = build_standard_cv(ROOT, self.work / "audited.pdf")
        audit = result["layout"]["audit"]
        self.assertEqual(audit["overflow_count"], 0)
        self.assertEqual(audit["violations"], [])
        self.assertGreaterEqual(len(audit["gaps"]), 5)
        self.assertEqual({round(item["actual"], 2) for item in audit["gaps"]}, {8.0})
        self.assertTrue(all(item["actual"] == item["expected"] for item in audit["gaps"]))
        self.assertIn(
            "agencement-sur-mesure-2020-2026:to-next-experience",
            {item["name"] for item in audit["gaps"]},
        )

    def test_entrepreneurial_cards_share_one_aligned_row(self) -> None:
        result = build_standard_cv(ROOT, self.work / "entrepreneurial-cards.pdf")
        frames = [
            frame
            for frame in result["layout"]["audit"]["frames"]
            if str(frame["name"]).startswith("highlight:Périmètre entrepreneurial:")
        ]
        self.assertEqual(len(frames), 3)
        self.assertEqual({float(frame["top"]) for frame in frames}, {float(frames[0]["top"])})
        self.assertEqual({float(frame["bottom"]) for frame in frames}, {float(frames[0]["bottom"])})

    def test_vertical_bars_match_their_content_boxes_and_rendered_paths(self) -> None:
        output = self.work / "bars.pdf"
        result = build_standard_cv(ROOT, output)
        audit = result["layout"]["audit"]
        elements = {item["name"]: item for item in audit["elements"]}
        bars = [item for item in audit["elements"] if item["kind"] == "vertical-bar"]
        self.assertEqual(len(bars), 9)
        for bar in bars:
            content = elements[bar["name"].replace(":bar", ":content")]
            self.assertEqual(bar["top"], content["top"])
            self.assertEqual(bar["bottom"], content["bottom"])
        self.assertTrue(all(item["actual"] == item["expected"] for item in audit["alignments"]))

        reader = PdfReader(str(output))
        stream = ContentStream(reader.pages[0].get_contents(), reader)
        path_x: list[float] = []
        path_y: list[float] = []
        rendered_bars: list[tuple[float, float]] = []
        for operands, operator in stream.operations:
            if operator in {b"m", b"l"}:
                path_x.append(float(operands[0]))
                path_y.append(float(operands[1]))
            elif operator == b"c":
                path_x.extend(float(operands[index]) for index in (0, 2, 4))
                path_y.extend(float(operands[index]) for index in (1, 3, 5))
            elif operator == b"re":
                origin_x = float(operands[0])
                origin_y = float(operands[1])
                path_x.extend((origin_x, origin_x + float(operands[2])))
                path_y.extend((origin_y, origin_y + float(operands[3])))
            elif operator in {b"B", b"B*", b"b", b"b*", b"S", b"s", b"f", b"f*", b"n"}:
                if (
                    path_x
                    and operator in {b"f", b"f*"}
                    and abs(min(path_x) - LEFT) < 0.01
                    and abs(max(path_x) - (LEFT + 3)) < 0.01
                    and max(path_y) - min(path_y) > 30
                ):
                    rendered_bars.append((min(path_y), max(path_y)))
                path_x = []
                path_y = []
        expected_page_one = sorted(
            (float(item["bottom"]), float(item["top"])) for item in bars if item["page"] == 1
        )
        self.assertEqual(len(rendered_bars), len(expected_page_one))
        for actual, expected in zip(sorted(rendered_bars), expected_page_one, strict=True):
            self.assertAlmostEqual(actual[0], expected[0], places=3)
            self.assertAlmostEqual(actual[1], expected[1], places=3)

    def test_rendered_card_strokes_stay_inside_safe_horizontal_margins(self) -> None:
        output = self.work / "paths.pdf"
        build_standard_cv(ROOT, output)
        reader = PdfReader(str(output))
        stream = ContentStream(reader.pages[0].get_contents(), reader)
        stroke_width = 1.0
        path_x: list[float] = []
        path_y: list[float] = []
        frame_bottoms: list[float] = []
        audited_frames = 0
        path_operators = {b"m", b"l", b"c", b"re"}
        paint_operators = {b"B", b"B*", b"b", b"b*", b"S", b"s", b"f", b"f*", b"n"}
        for operands, operator in stream.operations:
            if operator == b"w":
                stroke_width = float(operands[0])
            elif operator in path_operators:
                if operator in {b"m", b"l"}:
                    path_x.append(float(operands[0]))
                    path_y.append(float(operands[1]))
                elif operator == b"c":
                    path_x.extend(float(operands[index]) for index in (0, 2, 4))
                    path_y.extend(float(operands[index]) for index in (1, 3, 5))
                else:
                    origin_x = float(operands[0])
                    origin_y = float(operands[1])
                    path_x.extend((origin_x, origin_x + float(operands[2])))
                    path_y.extend((origin_y, origin_y + float(operands[3])))
            elif operator in paint_operators:
                if path_x and operator in {b"B", b"B*"} and abs(stroke_width - 0.8) < 0.01:
                    audited_frames += 1
                    self.assertGreaterEqual(min(path_x) - (stroke_width / 2), LEFT)
                    self.assertLessEqual(max(path_x) + (stroke_width / 2), RIGHT + 0.001)
                    frame_bottoms.append(min(path_y) - (stroke_width / 2))
                path_x = []
                path_y = []
        self.assertEqual(audited_frames, 7)

        date_runs: list[tuple[float, float]] = []

        def capture_date(text: str, _cm: list[float], tm: list[float], _font: object, size: float) -> None:
            if "03/2019 - 07/2020" in text:
                date_runs.append((float(tm[5]), float(size)))

        reader.pages[0].extract_text(visitor_text=capture_date)
        self.assertEqual(len(date_runs), 1)
        date_baseline, date_size = date_runs[0]
        date_top = date_baseline + getAscent("Helvetica-Bold", date_size)
        common_foundation_bottom = min(frame_bottoms)
        self.assertAlmostEqual(common_foundation_bottom - date_top, 8.0, delta=0.05)


if __name__ == "__main__":
    unittest.main()
